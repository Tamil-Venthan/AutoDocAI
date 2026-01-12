import os
import sys
import json
import re
import time
import threading
import logging
import queue
import webbrowser
import hashlib
import tkinter as tk
from datetime import timedelta
from tkinter import filedialog, messagebox, scrolledtext, simpledialog
from docx import Document
from docx.shared import Pt, RGBColor
from dotenv import load_dotenv
from openai import OpenAI, APIError, RateLimitError, APITimeoutError
import ttkbootstrap as ttk
from ttkbootstrap.constants import *

# ================= CONFIGURATION =================
VERSION = "v1.0.0"
SETTINGS_FILE = "settings.json"
TEMPLATES_FILE = "templates.json"
LOG_FILE = "run.log"
DEBUG_MODE = False

MIN_QUESTION_LENGTH = 15
MAX_TOKENS = 1500
USD_TO_INR = 86.0 

# Social Links
LINKEDIN_URL = "https://www.linkedin.com/in/tamil-venthan4"
GITHUB_URL = "https://github.com/Tamil-Venthan"
UPDATE_URL = "https://github.com/Tamil-Venthan/AutoDocAI/releases" 

MODEL_PRICING = {
    "gpt-4o": (2.50, 10.00),
    "gpt-4o-mini": (0.15, 0.60),
    "gpt-3.5-turbo": (0.50, 1.50)
}

DEFAULT_TEMPLATES = {
    "UPSC Mains Expert": "You are a UPSC Civil Services expert. Answer strictly in UPSC Mains format: Introduction, Body (with headings & bullets), and Conclusion.",
    "UPSC Ethics Expert": "You are an expert in UPSC Ethics. Provide answers with real-life examples, case studies, and ethical frameworks.",
    "UPSC GS Expert": "You are a UPSC General Studies expert. Provide detailed, well-structured answers with relevant data and examples.",
    "UPSC Essay Expert": "You are an expert essay writer for UPSC. Craft comprehensive essays with clear introductions, coherent arguments, and impactful conclusions.",
    "AI Tutor": "You are an AI tutor. Explain concepts clearly and concisely, using examples and analogies where appropriate.",
    "Tech Specialist": "You are a technology specialist. Provide detailed technical explanations and insights.",
    "Medical Expert": "You are a medical expert. Provide accurate and detailed medical information and explanations.",
    "Legal Advisor": "You are a legal advisor. Provide clear and precise legal explanations and advice.",
    "Financial Analyst": "You are a financial analyst. Provide detailed financial insights and analysis.",
    "Scientific Researcher": "You are a scientific researcher. Provide thorough and evidence-based scientific explanations.",
    "Creative Writer": "You are a creative writer. Craft engaging and imaginative content with vivid descriptions.",
    "General Professional": "You are a domain expert. Answer clearly and professionally."
}

USER_GUIDE_TEXT = """
📘 AutoDoc AI – User Guide
------------------------------------------
▶️ HOW TO USE:
1. Enter OpenAI API Key.
2. Select Input DOCX.
3. Select a Profile.
4. Click START.

🧠 Recommended AI Models

⭐ Best Overall (Recommended)
  Model: `gpt-4o`  
Best for:
- Mains answer writing
- Essay practice
- Ethics case studies
- Governance & policy analysis

Temperature:`0.2 – 0.4`


⚖️ Balanced Choice
Model: `gpt-4.1`  
Best for:
- GS Paper I–III
- Optional subjects
- Conceptual explanations

Temperature:`0.3 – 0.5`


💰 Budget & Fast
Model:`gpt-4o-mini`  
Best for:
- Daily revision
- Short notes
- MCQ explanations
- Large question banks

Temperature:`0.4 – 0.6`

💰 COST ESTIMATE (per 100 Qs):
• gpt-4o-mini: ₹10–₹20
• gpt-4o: ₹50–₹80
"""

# Setup Logging
log_level = logging.DEBUG if DEBUG_MODE else logging.INFO
logging.basicConfig(filename=LOG_FILE, level=log_level, format="%(asctime)s - %(levelname)s - %(message)s")
load_dotenv()

# ================= API WRAPPER =================
class OpenAIClient:
    def __init__(self, api_key):
        self.client = OpenAI(api_key=api_key)

    def generate_answer(self, system_prompt, user_prompt, model, temp, max_tokens):
        retries = 3
        base_delay = 2
        for attempt in range(retries):
            try:
                response = self.client.chat.completions.create(
                    model=model,
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt}
                    ],
                    temperature=temp,
                    max_tokens=max_tokens
                )
                return response
            except (RateLimitError, APITimeoutError) as e:
                wait_time = base_delay * (2 ** attempt)
                logging.warning(f"Retry ({attempt+1}) due to: {e}. Waiting {wait_time}s...")
                time.sleep(wait_time)
            except Exception as e:
                logging.exception("API Error")
                raise e
        raise Exception("Max retries exceeded.")

# ================= HELPER FUNCTIONS =================
def get_progress_filename(input_path):
    file_hash = hashlib.md5(input_path.encode('utf-8')).hexdigest()
    return f"progress_{file_hash}.json"

def load_json(file, default):
    if os.path.exists(file):
        try:
            with open(file, "r") as f: return json.load(f)
        except: pass
    return default

def save_json(file, data):
    try:
        with open(file, "w") as f: json.dump(data, f)
    except: pass

# --- FIXED FORMATTING LOGIC START ---
def add_formatted_text(paragraph, text):
    """
    Splits text by '**' and applies bold formatting to the bold parts.
    Example: "This is **important** text" -> "This is " (normal) + "important" (bold) + " text" (normal)
    """
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            # Remove the asterisks and make bold
            clean_part = part[2:-2]
            paragraph.add_run(clean_part).bold = True
        else:
            paragraph.add_run(part)

def parse_markdown_to_docx(doc, text):
    lines = text.split('\n')
    in_code_block = False

    for line in lines:
        stripped = line.strip()
        
        # 1. Handle Code Blocks
        if stripped.startswith("```"):
            in_code_block = not in_code_block
            continue
        
        if in_code_block:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0, 100, 0) # Green code
            p.paragraph_format.left_indent = Pt(20)
            continue
        
        # 2. Handle Headings (Fixed ###)
        if stripped.startswith("# "):
            doc.add_paragraph(stripped[2:], style='Heading 1')
        elif stripped.startswith("## "):
            doc.add_paragraph(stripped[3:], style='Heading 2')
        elif stripped.startswith("### "):
            doc.add_paragraph(stripped[4:], style='Heading 3')
        elif stripped.startswith("#### "):
            doc.add_paragraph(stripped[5:], style='Heading 4')
            
        # 3. Handle Lists
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, stripped[2:]) # Apply formatting inside list
            
        elif re.match(r"^\d+\.", stripped):
            p = doc.add_paragraph(style='List Number')
            # Remove "1. " from start
            clean_text = re.sub(r"^\d+\.\s*", "", stripped)
            add_formatted_text(p, clean_text)

        # 4. Normal Text
        else:
            if not stripped: continue
            p = doc.add_paragraph()
            add_formatted_text(p, line)
# --- FIXED FORMATTING LOGIC END ---

# ================= MAIN APP =================
class AutoDocAI:
    def __init__(self, root):
        self.root = root
        self.root.title(f"AutoDoc AI {VERSION}")
        self.root.geometry("1100x1000")
        self.root.protocol("WM_DELETE_WINDOW", self.on_close)

        self.root.bind("<Control-o>", lambda e: self.select_input())
        self.root.bind("<Control-Return>", lambda e: self.start())
        self.root.bind("<Escape>", lambda e: self.stop())

        self.msg_queue = queue.Queue()
        self.stop_event = threading.Event()
        self.worker_thread = None
        self.is_paused = False
        self.questions = []
        
        self.templates = load_json(TEMPLATES_FILE, DEFAULT_TEMPLATES)
        self.settings = load_json(SETTINGS_FILE, {"theme": "cyborg", "last_template": "UPSC Mains Expert"})

        self.input_path = tk.StringVar()
        self.output_path = tk.StringVar()
        self.model_var = tk.StringVar(value="gpt-4o-mini")
        self.temp_var = tk.DoubleVar(value=0.5)
        self.api_key_var = tk.StringVar(value=os.getenv("OPENAI_API_KEY", ""))
        self.show_key = tk.BooleanVar(value=False)
        self.current_template = tk.StringVar(value=self.settings.get("last_template", "UPSC Mains Expert"))
        self.current_theme = tk.StringVar(value=self.settings.get("theme", "cyborg"))
        self.page_break_var = tk.BooleanVar(value=False)
        
        self.stats = {"cost": 0.0, "processed": 0}
        self.setup_ui()
        self.root.after(100, self.process_queue)

    def setup_ui(self):
        main_frame = ttk.Frame(self.root, padding=20)
        main_frame.pack(fill=BOTH, expand=True)

        # Header
        header = ttk.Frame(main_frame)
        header.pack(fill=X, pady=(0, 10))
        ttk.Label(header, text="AutoDoc AI", font=("Helvetica", 18, "bold"), bootstyle="info").pack(side=LEFT)
        theme_fr = ttk.Frame(header)
        theme_fr.pack(side=RIGHT)
        ttk.Label(theme_fr, text="Theme:").pack(side=LEFT)
        self.cb_theme = ttk.Combobox(theme_fr, textvariable=self.current_theme, values=["cyborg", "superhero", "darkly", "cosmo", "flatly"], state="readonly", width=10)
        self.cb_theme.pack(side=LEFT, padx=5)
        self.cb_theme.bind("<<ComboboxSelected>>", self.change_theme)

        # Auth
        auth_fr = ttk.Labelframe(main_frame, text=" 🔑 API Key ", padding=10, bootstyle="secondary")
        auth_fr.pack(fill=X, pady=5)
        self.ent_api = ttk.Entry(auth_fr, textvariable=self.api_key_var, show="•")
        self.ent_api.pack(side=LEFT, fill=X, expand=True, padx=(0,10))
        ttk.Checkbutton(auth_fr, text="Show", variable=self.show_key, bootstyle="round-toggle",
                        command=lambda: self.ent_api.config(show="" if self.show_key.get() else "•")).pack(side=LEFT)

        # Files
        file_fr = ttk.Labelframe(main_frame, text=" 📂 File Operations ", padding=10, bootstyle="primary")
        file_fr.pack(fill=X, pady=5)
        f_grid = ttk.Frame(file_fr)
        f_grid.pack(fill=X)
        ttk.Label(f_grid, text="Input File (Ctrl+O):").grid(row=0, column=0, sticky="w")
        self.ent_input = ttk.Entry(f_grid, textvariable=self.input_path)
        self.ent_input.grid(row=0, column=1, sticky="ew", padx=10)
        self.btn_browse = ttk.Button(f_grid, text="Browse", command=self.select_input)
        self.btn_browse.grid(row=0, column=2)
        f_grid.columnconfigure(1, weight=1)
        ttk.Checkbutton(file_fr, text="Insert Page Break after each Answer", variable=self.page_break_var, bootstyle="square-toggle").pack(anchor="w", pady=(5,0))

        # Prompt
        sys_fr = ttk.Labelframe(main_frame, text=" 🧠 System Instruction ", padding=10, bootstyle="warning")
        sys_fr.pack(fill=X, pady=5)
        tpl_bar = ttk.Frame(sys_fr)
        tpl_bar.pack(fill=X, pady=(0, 5))
        ttk.Label(tpl_bar, text="Profile:").pack(side=LEFT)
        self.cb_tpl = ttk.Combobox(tpl_bar, textvariable=self.current_template, values=list(self.templates.keys()), state="readonly", width=30)
        self.cb_tpl.pack(side=LEFT, padx=5)
        self.cb_tpl.bind("<<ComboboxSelected>>", self.load_template)
        
        ttk.Button(tpl_bar, text="💾 Save", command=self.save_new_template, bootstyle="outline-success", width=8).pack(side=RIGHT)
        ttk.Button(tpl_bar, text="🗑 Del", command=self.delete_template, bootstyle="outline-danger", width=6).pack(side=RIGHT, padx=5)
        
        self.txt_system = scrolledtext.ScrolledText(sys_fr, height=4, font=("Segoe UI", 10))
        self.txt_system.pack(fill=X)
        self.txt_system.insert(tk.END, self.templates.get(self.current_template.get(), ""))

        # Settings
        ctrl_fr = ttk.Frame(main_frame)
        ctrl_fr.pack(fill=X, pady=10)
        set_fr = ttk.Labelframe(ctrl_fr, text=" Parameters ", padding=10)
        set_fr.pack(side=LEFT, fill=BOTH, expand=True, padx=(0, 10))
        
        ttk.Label(set_fr, text="Model:").pack(side=LEFT)
        self.cb_model = ttk.Combobox(set_fr, textvariable=self.model_var, values=list(MODEL_PRICING.keys()), state="readonly", width=15)
        self.cb_model.pack(side=LEFT, padx=5)
        
        ttk.Label(set_fr, text="Temp:").pack(side=LEFT, padx=(10,5))
        self.slider_temp = ttk.Scale(set_fr, variable=self.temp_var, from_=0.0, to=1.0, command=self.update_temp_label)
        self.slider_temp.pack(side=LEFT, fill=X, expand=True)
        self.lbl_temp_val = ttk.Label(set_fr, text="0.5", width=4, anchor="center")
        self.lbl_temp_val.pack(side=LEFT, padx=(5,0))

        # Stats
        cost_fr = ttk.Labelframe(ctrl_fr, text=" Live Stats ", padding=10, bootstyle="success")
        cost_fr.pack(side=LEFT, fill=BOTH)
        self.lbl_cost = ttk.Label(cost_fr, text="$0.00 | ₹0.00", font=("Consolas", 11, "bold"), bootstyle="success")
        self.lbl_cost.pack(anchor="w")
        self.lbl_eta = ttk.Label(cost_fr, text="ETA: --:--", font=("Consolas", 10), bootstyle="secondary")
        self.lbl_eta.pack(anchor="w")

        # Buttons
        btn_fr = ttk.Frame(main_frame)
        btn_fr.pack(fill=X, pady=5)
        self.btn_start = ttk.Button(btn_fr, text="▶ START (Ctrl+Enter)", command=self.start, bootstyle="success", width=20)
        self.btn_start.pack(side=LEFT, padx=5)
        self.btn_pause = ttk.Button(btn_fr, text="⏸ PAUSE", command=self.toggle_pause, bootstyle="warning", width=15, state="disabled")
        self.btn_pause.pack(side=LEFT, padx=5)
        self.btn_stop = ttk.Button(btn_fr, text="⏹ STOP (Esc)", command=self.stop, bootstyle="danger", width=15, state="disabled")
        self.btn_stop.pack(side=LEFT, padx=5)

        # Logs
        self.progress = ttk.Floodgauge(main_frame, bootstyle="info", text="Ready", mask="{}%")
        self.progress.pack(fill=X, pady=10)
        log_fr = ttk.Labelframe(main_frame, text=" Logs ", padding=10)
        log_fr.pack(fill=BOTH, expand=True)
        self.log_box = scrolledtext.ScrolledText(log_fr, height=5, state='disabled', font=("Consolas", 9))
        self.log_box.pack(fill=BOTH, expand=True)

        # Footer
        ftr = ttk.Frame(main_frame)
        ftr.pack(side=BOTTOM, fill=X, pady=(10,0))
        ttk.Button(ftr, text="🔄 Check Updates", command=lambda: webbrowser.open(UPDATE_URL), bootstyle="link").pack(side=LEFT)
        ttk.Label(ftr, text="|").pack(side=LEFT, padx=5)
        ttk.Button(ftr, text="📘 User Guide", command=self.show_guide, bootstyle="link-info").pack(side=LEFT)
        ttk.Button(ftr, text="LinkedIn", command=lambda: webbrowser.open(LINKEDIN_URL), bootstyle="link-primary").pack(side=RIGHT)
        ttk.Button(ftr, text="GitHub", command=lambda: webbrowser.open(GITHUB_URL), bootstyle="link-secondary").pack(side=RIGHT)

    # ================= UI LOGIC =================
    def update_temp_label(self, val):
        try: self.lbl_temp_val.config(text=f"{float(val):.1f}")
        except: pass

    def show_guide(self):
        top = tk.Toplevel(self.root)
        top.title("User Guide")
        top.geometry("700x600")
        st = scrolledtext.ScrolledText(top, font=("Consolas", 10), padx=10, pady=10)
        st.pack(fill=BOTH, expand=True)
        st.insert(tk.END, USER_GUIDE_TEXT)
        st.config(state='disabled')

    def toggle_controls(self, enable=True):
        state = "normal" if enable else "disabled"
        read_only = "readonly" if enable else "disabled"
        self.ent_input.config(state=state)
        self.btn_browse.config(state=state)
        self.cb_tpl.config(state=read_only)
        self.txt_system.config(state=state)
        self.cb_model.config(state=read_only)
        self.slider_temp.config(state=state)
        self.cb_theme.config(state=read_only)
    
    def load_template(self, event):
        name = self.current_template.get()
        self.txt_system.delete("1.0", tk.END)
        self.txt_system.insert(tk.END, self.templates.get(name, ""))
        save_json(SETTINGS_FILE, {"last_template": name, "theme": self.current_theme.get()})

    def save_new_template(self):
        name = simpledialog.askstring("Save Profile", "Profile Name:")
        if name:
            self.templates[name] = self.txt_system.get("1.0", tk.END).strip()
            save_json(TEMPLATES_FILE, self.templates)
            self.cb_tpl['values'] = list(self.templates.keys())
            self.current_template.set(name)

    def delete_template(self):
        name = self.current_template.get()
        if name in self.templates and messagebox.askyesno("Confirm", f"Delete '{name}'?"):
            del self.templates[name]
            save_json(TEMPLATES_FILE, self.templates)
            self.cb_tpl['values'] = list(self.templates.keys())
            self.current_template.set(list(self.templates.keys())[0])
            self.load_template(None)

    def change_theme(self, event):
        t = self.current_theme.get()
        ttk.Style().theme_use(t)
        save_json(SETTINGS_FILE, {"last_template": self.current_template.get(), "theme": t})

    def select_input(self):
        p = filedialog.askopenfilename(filetypes=[("Word Files", "*.docx")])
        if p:
            self.input_path.set(p)
            d, f = os.path.split(p)
            self.output_path.set(os.path.join(d, f"{os.path.splitext(f)[0]}_Answers.docx"))
            if os.path.exists(p):
                doc = Document(p)
                self.questions = [para.text.strip() for para in doc.paragraphs if len(para.text.strip()) >= MIN_QUESTION_LENGTH]
                self.log_gui(f"Loaded {len(self.questions)} questions.")

    def log_gui(self, msg):
        self.msg_queue.put(("log", msg))

    def process_queue(self):
        try:
            while True:
                t, d = self.msg_queue.get_nowait()
                if t == "log":
                    self.log_box.config(state='normal')
                    self.log_box.insert(tk.END, f"• {d}\n")
                    self.log_box.see(tk.END)
                    self.log_box.config(state='disabled')
                elif t == "progress":
                    self.progress.configure(value=d['val'], text=d['text'])
                elif t == "stats":
                    self.lbl_cost.config(text=f"${d['usd']:.4f} | ₹{d['usd']*USD_TO_INR:.2f}")
                    self.lbl_eta.config(text=f"ETA: {d['eta']}")
                elif t == "done":
                    messagebox.showinfo("Done", "Processing Complete")
                    self.reset_ui()
        except queue.Empty: pass
        finally: self.root.after(100, self.process_queue)

    def toggle_pause(self):
        self.is_paused = not self.is_paused
        self.btn_pause.config(text="▶ RESUME" if self.is_paused else "⏸ PAUSE")
        self.log_gui("Paused" if self.is_paused else "Resumed")

    def stop(self):
        if self.worker_thread and self.worker_thread.is_alive():
            if messagebox.askyesno("Confirm", "Stop processing?"):
                self.stop_event.set()

    def on_close(self):
        if self.worker_thread and self.worker_thread.is_alive():
            if messagebox.askyesno("Exit", "Process running. Stop & Save?"):
                self.stop_event.set()
                self.root.attributes('-disabled', True) 
                self.root.after(200, self.check_safe_exit)
        else:
            self.root.destroy()

    def check_safe_exit(self):
        if self.worker_thread.is_alive():
            self.root.after(200, self.check_safe_exit)
        else:
            self.root.destroy()

    def reset_ui(self):
        self.toggle_controls(True)
        self.btn_start.config(state="normal")
        self.btn_pause.config(state="disabled")
        self.btn_stop.config(state="disabled")

    def start(self):
        if not self.questions: return messagebox.showwarning("Error", "Load file first.")
        api_key = self.api_key_var.get()
        if not api_key: return messagebox.showwarning("Error", "API Key required.")
        sys_prompt = self.txt_system.get("1.0", tk.END).strip()
        self.toggle_controls(False)
        self.btn_start.config(state="disabled")
        self.btn_pause.config(state="normal")
        self.btn_stop.config(state="normal")
        self.stop_event.clear()
        self.worker_thread = threading.Thread(target=self.worker, args=(sys_prompt, api_key), daemon=True)
        self.worker_thread.start()

    def worker(self, sys_prompt, api_key):
        input_path = self.input_path.get()
        in_file_name = os.path.basename(input_path)
        out_path = self.output_path.get()
        doc = Document(out_path) if os.path.exists(out_path) else Document()
        
        progress_file = get_progress_filename(in_file_name)
        start_idx = 0
        if os.path.exists(progress_file):
            data = load_json(progress_file, {})
            last_idx = data.get("last_index", -1)
            if last_idx >= 0 and last_idx < len(self.questions) - 1:
                start_idx = last_idx + 1
                self.log_gui(f"Resuming from Question {start_idx + 1}...")

        try:
            client = OpenAIClient(api_key)
            model = self.model_var.get()
            temp = self.temp_var.get()
            start_time = time.time()
            processed_count = 0

            for i in range(start_idx, len(self.questions)):
                if self.stop_event.is_set(): break
                while self.is_paused: time.sleep(0.5)

                q = self.questions[i]
                self.msg_queue.put(("progress", {"val": int((i/len(self.questions))*100), "text": f"Q{i+1}/{len(self.questions)}"}))

                try:
                    resp = client.generate_answer(sys_prompt, q, model, temp, MAX_TOKENS)
                    ans = resp.choices[0].message.content
                    pt = resp.usage.prompt_tokens
                    ct = resp.usage.completion_tokens
                    pi, po = MODEL_PRICING.get(model, (0,0))
                    cost = (pt/1e6 * pi) + (ct/1e6 * po)
                    processed_count += 1
                    self.stats['cost'] += cost
                    
                    elapsed = time.time() - start_time
                    avg_time = elapsed / processed_count
                    remaining_qs = len(self.questions) - (i + 1)
                    eta_str = str(timedelta(seconds=int(avg_time * remaining_qs)))

                    self.msg_queue.put(("stats", {
                        "usd": self.stats['cost'], 
                        "eta": eta_str
                    }))

                    p = doc.add_paragraph()
                    p.add_run(f"Q{i+1}: {q}").bold = True
                    p.style = 'Heading 2'
                    
                    # Call the FIXED format parser
                    parse_markdown_to_docx(doc, ans)
                    
                    doc.add_paragraph("_"*30)
                    if self.page_break_var.get(): doc.add_page_break()

                    save_json(progress_file, {"last_index": i})
                    try: doc.save(out_path)
                    except: self.log_gui("⚠️ Save delayed (File open)")

                except Exception as e:
                    self.log_gui(f"Error on Q{i+1}: {e}")

            self.msg_queue.put(("progress", {"val": 100, "text": "Finished"}))
            self.msg_queue.put(("done", None))
            if os.path.exists(progress_file): os.remove(progress_file)

        except Exception as e:
            self.log_gui(f"Critical Worker Error: {e}")
        finally:
            try: doc.save(out_path)
            except: pass

if __name__ == "__main__":
    s = load_json(SETTINGS_FILE, {"theme": "cyborg"})
    app = ttk.Window(themename=s.get("theme", "cyborg")) 
    AutoDocAI(app)
    app.mainloop()